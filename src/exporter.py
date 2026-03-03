"""
Multi-format exporter for State of the Art documents.

Supports LaTeX, Markdown, and Word (docx) output formats.
"""

import re
from pathlib import Path
from typing import Optional, TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx import Document as DocxDocument

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = Any  # type: ignore


class SOAExporter:
    """
    Export State of the Art documents to multiple formats.
    
    Supports:
    - LaTeX (.tex)
    - Markdown (.md)
    - Word (.docx)
    """
    
    def __init__(self):
        """Initialize exporter."""
        pass
    
    def to_latex(self, content: str, output_path: str) -> str:
        """
        Export to LaTeX format.
        
        Args:
            content: LaTeX content (already in LaTeX format)
            output_path: Output file path
            
        Returns:
            Output file path
        """
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # LaTeX content is typically already in correct format
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"  ✓ LaTeX exported: {output_file}")
        return str(output_file)
    
    def to_markdown(self, content: str, output_path: str) -> str:
        """
        Convert LaTeX to Markdown and export.
        
        Args:
            content: LaTeX content
            output_path: Output file path
            
        Returns:
            Output file path
        """
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Convert LaTeX to Markdown
        markdown = self._latex_to_markdown(content)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown)
        
        print(f"  ✓ Markdown exported: {output_file}")
        return str(output_file)
    
    def to_docx(self, content: str, output_path: str) -> str:
        """
        Convert LaTeX to Word document and export.
        
        Args:
            content: LaTeX content
            output_path: Output file path
            
        Returns:
            Output file path
            
        Raises:
            ImportError: If python-docx is not installed
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word export. "
                "Install it with: pip install python-docx"
            )
        
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Create Word document
        doc = Document()  # type: ignore
        
        # Parse LaTeX and populate document
        self._latex_to_docx(content, doc)
        
        # Save document
        doc.save(str(output_file))
        
        print(f"  ✓ Word document exported: {output_file}")
        return str(output_file)
    
    def _latex_to_markdown(self, latex_content: str) -> str:
        """
        Convert LaTeX to Markdown.
        
        Args:
            latex_content: LaTeX document content
            
        Returns:
            Markdown content
        """
        md = latex_content
        
        # Remove LaTeX document structure
        md = re.sub(r'\\documentclass\{.*?\}', '', md)
        md = re.sub(r'\\usepackage(\[.*?\])?\{.*?\}', '', md)
        md = re.sub(r'\\begin\{document\}', '', md)
        md = re.sub(r'\\end\{document\}', '', md)
        md = re.sub(r'\\maketitle', '', md)
        
        # Convert sections
        md = re.sub(r'\\section\{(.*?)\}', r'## \1', md)
        md = re.sub(r'\\subsection\{(.*?)\}', r'### \1', md)
        md = re.sub(r'\\subsubsection\{(.*?)\}', r'#### \1', md)
        
        # Convert title and author
        md = re.sub(r'\\title\{(.*?)\}', r'# \1', md)
        md = re.sub(r'\\author\{(.*?)\}', r'**Author:** \1', md)
        md = re.sub(r'\\date\{(.*?)\}', r'**Date:** \1', md)
        
        # Convert text formatting
        md = re.sub(r'\\textbf\{(.*?)\}', r'**\1**', md)
        md = re.sub(r'\\textit\{(.*?)\}', r'*\1*', md)
        md = re.sub(r'\\emph\{(.*?)\}', r'*\1*', md)
        md = re.sub(r'\\texttt\{(.*?)\}', r'`\1`', md)
        
        # Convert citations
        md = re.sub(r'\\cite\{(.*?)\}', r'[\1]', md)
        md = re.sub(r'\\citep\{(.*?)\}', r'[\1]', md)
        md = re.sub(r'\\citet\{(.*?)\}', r'[\1]', md)
        
        # Convert lists
        md = re.sub(r'\\begin\{itemize\}', '', md)
        md = re.sub(r'\\end\{itemize\}', '', md)
        md = re.sub(r'\\begin\{enumerate\}', '', md)
        md = re.sub(r'\\end\{enumerate\}', '', md)
        md = re.sub(r'\\item\s+', '- ', md)
        
        # Convert abstract
        md = re.sub(r'\\begin\{abstract\}', '## Abstract\n', md)
        md = re.sub(r'\\end\{abstract\}', '', md)
        
        # Remove other LaTeX commands
        md = re.sub(r'\\[a-zA-Z]+\{(.*?)\}', r'\1', md)
        md = re.sub(r'\\[a-zA-Z]+', '', md)
        
        # Clean up whitespace
        md = re.sub(r'\n{3,}', '\n\n', md)
        md = md.strip()
        
        return md
    
    def _latex_to_docx(self, latex_content: str, doc):
        """
        Parse LaTeX and populate Word document.
        
        Args:
            latex_content: LaTeX document content
            doc: python-docx Document object
        """
        # Extract title
        title_match = re.search(r'\\title\{(.*?)\}', latex_content)
        if title_match:
            title = title_match.group(1)
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Extract author
        author_match = re.search(r'\\author\{(.*?)\}', latex_content)
        if author_match:
            author = author_match.group(1)
            author_para = doc.add_paragraph(author)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Extract abstract
        abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', latex_content, re.DOTALL)
        if abstract_match:
            doc.add_heading('Abstract', level=1)
            abstract_text = abstract_match.group(1).strip()
            abstract_text = self._clean_latex_text(abstract_text)
            doc.add_paragraph(abstract_text)
        
        # Extract sections
        # Find all sections
        section_pattern = r'\\section\{(.*?)\}(.*?)(?=\\section\{|\\end\{document\}|$)'
        sections = re.findall(section_pattern, latex_content, re.DOTALL)
        
        for section_title, section_content in sections:
            # Add section heading
            doc.add_heading(section_title, level=1)
            
            # Parse section content
            # Find subsections
            subsection_pattern = r'\\subsection\{(.*?)\}(.*?)(?=\\subsection\{|\\section\{|$)'
            subsections = re.findall(subsection_pattern, section_content, re.DOTALL)
            
            if subsections:
                # Has subsections
                for subsection_title, subsection_content in subsections:
                    doc.add_heading(subsection_title, level=2)
                    text = self._clean_latex_text(subsection_content)
                    if text.strip():
                        # Parse paragraphs and handle citations
                        self._add_formatted_text(doc, text)
            else:
                # No subsections, add content directly
                text = self._clean_latex_text(section_content)
                if text.strip():
                    self._add_formatted_text(doc, text)
    
    def _clean_latex_text(self, text: str) -> str:
        """
        Clean LaTeX commands from text.
        
        Args:
            text: LaTeX text
            
        Returns:
            Cleaned text
        """
        # Remove comments
        text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
        
        # Convert citations (keep before general cleanup)
        text = re.sub(r'\\cite\{(.*?)\}', r'[\1]', text)
        text = re.sub(r'\\citep\{(.*?)\}', r'[\1]', text)
        text = re.sub(r'\\citet\{(.*?)\}', r'[\1]', text)
        
        # Convert text formatting
        text = re.sub(r'\\textbf\{(.*?)\}', r'**\1**', text)
        text = re.sub(r'\\textit\{(.*?)\}', r'*\1*', text)
        text = re.sub(r'\\emph\{(.*?)\}', r'*\1*', text)
        
        # Remove other LaTeX commands
        text = re.sub(r'\\[a-zA-Z]+\{(.*?)\}', r'\1', text)
        text = re.sub(r'\\[a-zA-Z]+', '', text)
        
        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()
        
        return text
    
    def _add_formatted_text(self, doc, text: str):
        """
        Add formatted text to document, handling inline formatting.
        
        Args:
            doc: python-docx Document object
            text: Text with markdown-style formatting
        """
        # Split into paragraphs
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            para = doc.add_paragraph()
            
            # Parse inline formatting
            # Simple approach: split by formatting markers and apply styles
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\[.*?\])', para_text)
            
            for part in parts:
                if not part:
                    continue
                
                if part.startswith('**') and part.endswith('**'):
                    # Bold
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    # Italic
                    run = para.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('[') and part.endswith(']'):
                    # Citation
                    run = para.add_run(part)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for citations
                else:
                    # Normal text
                    para.add_run(part)


def export_all_formats(content: str, output_dir: str = "artifacts/soa", 
                       basename: str = "state_of_the_art_final") -> dict:
    """
    Export to all formats (LaTeX, Markdown, Word).
    
    Args:
        content: LaTeX content
        output_dir: Output directory
        basename: Base filename (without extension)
        
    Returns:
        Dictionary mapping format names to file paths
    """
    exporter = SOAExporter()
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    results = {}
    
    # LaTeX
    try:
        latex_file = output_path / f"{basename}.tex"
        exporter.to_latex(content, str(latex_file))
        results['latex'] = str(latex_file)
    except Exception as e:
        print(f"  ✗ LaTeX export failed: {e}")
        results['latex'] = None
    
    # Markdown
    try:
        md_file = output_path / f"{basename}.md"
        exporter.to_markdown(content, str(md_file))
        results['markdown'] = str(md_file)
    except Exception as e:
        print(f"  ✗ Markdown export failed: {e}")
        results['markdown'] = None
    
    # Word
    try:
        docx_file = output_path / f"{basename}.docx"
        exporter.to_docx(content, str(docx_file))
        results['docx'] = str(docx_file)
    except Exception as e:
        print(f"  ✗ Word export failed: {e}")
        results['docx'] = None
    
    return results
